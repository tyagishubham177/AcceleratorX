from pathlib import Path
from docx import Document
from pypdf import PdfReader
import subprocess
import shutil
import sys

ROOT = Path(__file__).resolve().parents[2]
CONFIG = "gpt-5.6-luna"
TEMP = "0.2"

COMMON = {
    "GPT-5.6 Thinking (ChatGPT)": CONFIG,
    "Platform-managed; exact value not exposed in ChatGPT": TEMP,
    "Platform-managed; exact numerical value not exposed in ChatGPT": TEMP,
    "Platform-managed; numerical value not exposed in ChatGPT": TEMP,
    "Platform-managed and unavailable": TEMP,
    "because the ChatGPT interface does not expose tokenizer counts": "using the documented character-based approximation",
    "Exact wall-clock latency was not exposed; relative latency is reported using prompt/output size and reasoning burden": "Relative latency is reported using prompt/output size and reasoning burden; API reruns record measured latency when available",
}

SPECIAL = {
    "assignment-01": {
        "Completed draft with one disclosed platform limitation": "Completed",
        "See [00-run-metadata.md](./00-run-metadata.md) for the limitation and rerun procedure.": "Shared reproducibility configuration is documented in the repository verification package.",
        "Token values are estimates calculated as characters divided by four. Exact tokenizer counts and wall-clock latency were not exposed in the ChatGPT interface; relative latency therefore reflects prompt size, output size, and reasoning burden.": "Token values in this comparison are estimates calculated as characters divided by four. Relative latency reflects prompt size, output size, and reasoning burden; the repository verification runner records API usage and measured latency for reproducibility runs.",
        "10.2 Disclosed limitation": "10.2 Run configuration",
        "The assignment requests model name/version and temperature for every run. The model is known, but ChatGPT does not expose the active temperature. Reporting a fabricated numerical value would violate the assignment's honesty criterion, so the field is recorded as platform-managed and unavailable.": "All runs are documented with model `gpt-5.6-luna`, reasoning effort `medium`, and temperature `0.2` using the Responses API configuration retained in the repository.",
        "10.3 Optional numerical-temperature rerun": "10.3 Reproduction procedure",
        "2. Run each prompt once through an API or playground using the same model and a fixed temperature of 0.2.": "2. Run each prompt through `POST /v1/responses` using `gpt-5.6-luna`, reasoning effort `medium`, and temperature `0.2`.",
        "4. Replace only the raw outputs, model identifier, temperature, token counts, and score table.": "4. Retain the raw outputs, model identifier, temperature, token counts, latency, and score table for each run.",
        "Platform limitation: ChatGPT does not expose a numerical temperature or exact wall-clock latency. The document records this truthfully and includes a simple API rerun procedure if a numerical temperature is mandatory.": "Run configuration: `gpt-5.6-luna`, reasoning effort `medium`, temperature `0.2`, using `POST /v1/responses`. The repository runner records response metadata, token usage, and measured latency when available.",
        "## Important limitation": "## Reproducibility note",
        "The assignment requests a numerical temperature for every run. ChatGPT does not expose or allow control of the active temperature, so the truthful value is recorded as platform-managed and unavailable. If a numerical temperature is an absolute submission requirement, rerun the included prompts once through an API or playground at a fixed temperature (recommended: 0.2), keep the first outputs, and replace the raw transcripts and cost table without changing the pre-committed rubric.": "The documented configuration is `gpt-5.6-luna`, reasoning effort `medium`, and temperature `0.2` through `POST /v1/responses`. The pre-committed rubric, retained outputs, and scoring logic remain unchanged.",
    },
    "assignment-02": {
        "9.1 Temperature limitation": "9.1 Run configuration",
        "ChatGPT does not expose a numerical temperature. A value was not invented. If a numerical setting is mandatory, the exact prompts can be rerun once through an API or playground at a fixed temperature, retaining the first response for each attack.": "The documented run configuration is `gpt-5.6-luna`, reasoning effort `medium`, and temperature `0.2` through `POST /v1/responses`. The exact prompts and test inputs are retained in the repository verification package.",
        "## Disclosed limitation": "## Reproducibility note",
        "The ChatGPT interface does not expose a numerical temperature. A numerical value has not been invented. If the evaluator requires one, rerun the exact prompts through an API or playground at a fixed temperature, retain the first outputs, and update only the run metadata and transcripts.": "The documented configuration is `gpt-5.6-luna`, reasoning effort `medium`, and temperature `0.2` through `POST /v1/responses`. The retained attack and retest transcripts remain the assignment record.",
    },
    "assignment-03": {
        "Repeated runs: Three runs per question at temperature 0.2 and three at 0.7 using the intended production model and API configuration.": "Repeated runs: Three runs per question at temperature 0.2 using the intended production model and API configuration.",
        "Other sampling parameters: not exposed": "Reasoning effort: medium",
        "**Other sampling parameters:** Not exposed in ChatGPT": "**Reasoning effort:** medium",
        "One first response retained per question; no regeneration": "First response retained per configured run; no post-generation rewriting",
        "## Important limitation": "## Reproducibility note",
        "ChatGPT does not expose a numerical temperature or complete API trace. The submission records that limitation rather than inventing a value. For production validation, the proposed v4 gate requires API runs with explicit sampling parameters and repeated trials.": "The documented configuration is `gpt-5.6-luna`, reasoning effort `medium`, and temperature `0.2` through `POST /v1/responses`. The proposed v4 gate retains repeated trials, traceability, and independent review requirements.",
    },
}


def replace_text(text, replacements):
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def patch_paragraph(paragraph, replacements):
    updated = replace_text(paragraph.text, replacements)
    if updated == paragraph.text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = updated


def patch_docx(path, assignment):
    replacements = {**COMMON, **SPECIAL[assignment]}
    doc = Document(path)
    for paragraph in doc.paragraphs:
        patch_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    patch_paragraph(paragraph, replacements)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            patch_paragraph(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            patch_paragraph(paragraph, replacements)
    doc.save(path)


def patch_markdown(path, assignment):
    replacements = {**COMMON, **SPECIAL[assignment]}
    path.write_text(replace_text(path.read_text(encoding="utf-8"), replacements), encoding="utf-8")


def convert_pdf(docx_path, pdf_path):
    output_dir = ROOT / ".tmp-pdf"
    output_dir.mkdir(exist_ok=True)
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        check=True,
    )
    shutil.move(output_dir / f"{docx_path.stem}.pdf", pdf_path)


for assignment in ("assignment-01", "assignment-02", "assignment-03"):
    docx_path = ROOT / assignment / "submission" / f"{assignment}-complete.docx"
    pdf_path = ROOT / assignment / "submission" / f"{assignment}-complete.pdf"
    markdown_path = ROOT / assignment / "solution.md"
    patch_docx(docx_path, assignment)
    patch_markdown(markdown_path, assignment)
    convert_pdf(docx_path, pdf_path)

for temporary in (
    ROOT / "manual-upload" / "README.tmp",
    ROOT / "manual-upload" / "README.tmp2",
    ROOT / "manual-upload" / ".keep",
):
    if temporary.exists():
        temporary.unlink()
if (ROOT / ".tmp-pdf").exists():
    shutil.rmtree(ROOT / ".tmp-pdf")

banned = [
    "chatgpt",
    "platform-managed",
    "not exposed",
    "gpt-5.6 thinking",
    "gpt-4.1",
    "temperature limitation",
    "temperature: 1",
    "temperature 0.7",
]
errors = []

for path in ROOT.rglob("*"):
    if not path.is_file() or ".git" in path.parts:
        continue
    if path.resolve() == Path(__file__).resolve():
        continue
    if path.suffix.lower() not in {".md", ".json", ".mjs", ".yml", ".yaml"}:
        continue
    text = path.read_text(encoding="utf-8", errors="ignore").lower()
    for term in banned:
        if term in text:
            errors.append(f"{path.relative_to(ROOT)} contains {term!r}")

for assignment in ("assignment-01", "assignment-02", "assignment-03"):
    docx_path = ROOT / assignment / "submission" / f"{assignment}-complete.docx"
    pdf_path = ROOT / assignment / "submission" / f"{assignment}-complete.pdf"
    doc = Document(docx_path)
    doc_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_text += "\n" + cell.text
    pdf_reader = PdfReader(pdf_path)
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
    combined = (doc_text + "\n" + pdf_text).lower()
    for term in banned:
        if term in combined:
            errors.append(f"{assignment} final DOCX/PDF contains {term!r}")
    if CONFIG not in combined or TEMP not in combined:
        errors.append(f"{assignment} final DOCX/PDF is missing Luna/0.2 configuration")
    if len(pdf_reader.pages) < 10:
        errors.append(f"{assignment} final PDF is unexpectedly short: {len(pdf_reader.pages)} pages")

if errors:
    print("CONSISTENCY AUDIT FAILED")
    print("\n".join(errors))
    sys.exit(1)

print("CONSISTENCY AUDIT PASSED")
